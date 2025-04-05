import os
import docx
import PyPDF2
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain_core.prompts.prompt import PromptTemplate
from langchain import LLMChain
import streamlit as st
from langchain.schema import Document

def read_document(uploaded_file, file_type):
    try:
        if file_type == ".docx":
            # Pass the file-like object to docx.Document.
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file_type == ".txt":
            # Ensure the pointer is at the beginning and decode bytes to string.
            uploaded_file.seek(0)
            return uploaded_file.read().decode('utf-8').strip()
        elif file_type == ".pdf":
            # Use the uploaded file directly with PyPDF2.
            reader = PyPDF2.PdfReader(uploaded_file)
            return "\n".join([page.extract_text() for page in reader.pages])
        else:
            st.warning("Unsupported file type: %s" % file_type)
            return "Unsupported file type. Please upload a .docx, .txt, or .pdf file."
    except Exception as e:
        logging.error("Error reading file: %s", e)
        return f"Error reading file: {e}"

def get_info(resume_text, openai_api_key, open_ai_model):
    document = Document(page_content=resume_text, metadata={})
    doc_splitter_recursive = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = doc_splitter_recursive.split_documents([document])

    embed_object = OpenAIEmbeddings(openai_api_key=openai_api_key)
    vectordb = FAISS.from_documents(documents=split_docs, embedding=embed_object)

    prompt_template = """Answer the following question using only the resume content provided below.
                        If you do not know the answer, do not attempt to guess.
                        Use only the information given, without adding any extra details.
                        Keep the answers as minimal as possible

                        Resume: {context}

                        Question: {question}"""

    PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain_type_kwargs = {"prompt": PROMPT}
    retriever = vectordb.as_retriever(search_type="similarity", search_kwargs={"k": 10})
    llm = ChatOpenAI(model_name=open_ai_model, temperature=0.6, openai_api_key=openai_api_key)
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever, chain_type_kwargs=chain_type_kwargs)
    return qa_chain

def extract_qualifications(job_content, openai_api_key, open_ai_model):
    document = Document(page_content=job_content, metadata={})
    doc_splitter_recursive = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = doc_splitter_recursive.split_documents([document])

    embed_object = OpenAIEmbeddings(openai_api_key=openai_api_key)
    vectordb = FAISS.from_documents(documents=split_docs, embedding=embed_object)
    
    prompt_template = PromptTemplate(
        input_variables=["context"],
        template=(
            "You are an expert at extracting structured information from job descriptions. "
            "From the given job description, extract the qualifications section in full detail. "
            "Make sure to include all qualifications, tools, technologies, and specific examples mentioned. "
            "Format your response as a list with each qualification.\n\n"
            "Job Description:\n{context}\n\n"
            "Qualifications:\n"
        )
    )
    chain_type_kwargs = {"prompt": prompt_template}
    retriever = vectordb.as_retriever(search_type="similarity", search_kwargs={"k": 5})
    llm = ChatOpenAI(model_name=open_ai_model, temperature=0, openai_api_key=openai_api_key)
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever, chain_type_kwargs=chain_type_kwargs, return_source_documents=True)
    qualifications = qa_chain.invoke({"query": "Extract qualifications from this job description."})["result"]
    return qualifications.strip().split("\n")

def check_resume_against_qualifications(qa_chain, qualifications):
    results = {}
    for qual in qualifications:
        question = f"Does the resume meet the following qualification? {qual}"
        response = qa_chain.invoke({"query": question})["result"]
        results[qual] = response
    return results

def main():
    st.title("Resume Qualification Checker")
    openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")
    open_ai_model = st.text_input("Enter OpenAI Model Name", value="gpt-4o")

    # File uploaders
    job_description_file = st.file_uploader("Upload Job Description (DOCX, PDF, or TXT)", type=["docx", "pdf", "txt"])
    resume_file = st.file_uploader("Upload Resume (DOCX, PDF, or TXT)", type=["docx", "pdf", "txt"])

    # Ensure session state for process and question history
    if "processed" not in st.session_state:
        st.session_state.processed = False
    if "qa_history" not in st.session_state:
        st.session_state.qa_history = []

    # UI Change: Added a Start Process button to initiate processing
    if openai_api_key and job_description_file and resume_file and st.button("Start Process"):
        # Read job description
        job_ext = os.path.splitext(job_description_file.name)[1]
        job_content = read_document(job_description_file, job_ext)

        # Extract qualifications from job description
        st.info("Extracting qualifications from the job description...")
        qualifications = extract_qualifications(job_content, openai_api_key, open_ai_model)
        st.success("Qualifications extracted!")

        # Read resume
        resume_ext = os.path.splitext(resume_file.name)[1]
        resume_text = read_document(resume_file, resume_ext)

        # Build QA chain for resume
        st.info("Building QA chain for the resume...")
        qa_chain = get_info(resume_text, openai_api_key, open_ai_model)
        st.success("QA chain built!")

        # Check resume against qualifications
        st.info("Checking resume against qualifications...")
        results = check_resume_against_qualifications(qa_chain, qualifications)
        st.success("Completed checking!")

        # Save results in session state for later use
        st.session_state.qa_chain = qa_chain
        st.session_state.qualifications = qualifications
        st.session_state.results = results
        st.session_state.processed = True

    # Display processed results and the question wall if processing is complete
    if st.session_state.processed:
        col1, col2 = st.columns([3, 2], gap="medium")
        
        with col1:
            #st.write("### Extracted Qualifications:")
            #for q in st.session_state.qualifications:
            #    st.write(f"- {q}")
    
            st.write("### Results:")
            for qualification, result in st.session_state.results.items():
                st.markdown(f"**Qualification:** {qualification}")
                st.markdown(f"**Result:** {result}")
                st.write("---")

            # UI Change: Added question input and history wall
            st.write("### Ask Questions About the Resume:")
            custom_question = st.text_input(
                "Ask a question about the resume (e.g., 'What is the candidate's most recent job role?')",
                value="",
                key="custom_question"
            )

        with col2:
            if st.button("Ask Question"):
                if custom_question.strip():
                    st.info("Generating answer...")
                    response = st.session_state.qa_chain.invoke({"query": custom_question})["result"]
                    # Append the question and answer to the question history wall
                    st.session_state.qa_history.append({"question": custom_question, "answer": response})
                    st.success("Question added to the wall!")
    
            # Display the question history wall
            st.write("### Question History:")
            for item in st.session_state.qa_history:
                st.markdown(f"**Question:** {item['question']}")
                st.markdown(f"**Answer:** {item['answer']}")
                st.write("---")
    else:
        st.error("Please provide your OpenAI API key, a job description, and a resume to continue.")

if __name__ == "__main__":
    main()
